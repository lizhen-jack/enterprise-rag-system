"""
文档解析器
支持PDF、Word、Excel、纯文本
"""

from typing import List
import os


async def parse_pdf(file_path: str) -> str:
    """解析PDF文档"""
    import pypdf

    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = pypdf.PdfReader(file)

        # 提取所有页面文本
        for page in pdf_reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"

    return text.strip()


async def parse_docx(file_path: str) -> str:
    """解析Word文档"""
    from docx import Document

    doc = Document(file_path)
    text = ""

    # 提取段落文本
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    # 提取表格文本
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text for cell in row.cells])
            if row_text.strip():
                text += row_text + "\n"

    return text.strip()


async def parse_text(file_path: str) -> str:
    """解析纯文本文件"""
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()

    return text.strip()


async def parse_excel(file_path: str) -> str:
    """解析Excel文件"""
    from openpyxl import load_workbook

    wb = load_workbook(file_path)
    text = ""

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text += f"工作表: {sheet_name}\n"

        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
            if row_text.strip():
                text += row_text + "\n"

    return text.strip()
