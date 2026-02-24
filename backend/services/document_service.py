"""
文档处理服务
支持：PDF/DOCX/TXT/MD/XLSX解析 + 自动分块 + 向量索引
"""

import os
import hashlib
import fitz  # PyMuPDF
from docx import Document
from openpyxl import load_workbook
from typing import List, Dict, Any, Optional
from datetime import datetime
import aiofiles
import aiofiles.os as aios

from models.document import Document
from sqlalchemy import select
from services.rag_service import RAGService


class DocumentService:
    """文档处理服务"""

    def __init__(self, db, rag_service: RAGService):
        self.db = db
        self.rag_service = rag_service
        self.upload_dir = settings.UPLOAD_DIR
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP

    async def upload_document(
        self,
        user_id: int,
        file_path: str,
        filename: str,
        file_size: int,
        mime_type: str
    ) -> Document:
        """上传文档"""
        # 计算文件哈希
        file_hash = await self._calculate_file_hash(file_path)

        # 检查是否已存在
        query = select(Document).where(
            Document.user_id == user_id,
            Document.file_hash == file_hash
        )
        result = await self.db.execute(query)
        existing = result.scalar_one_or_none()

        if existing:
            return existing

        # 创建文档记录
        # 从filename中提取文件名（不带扩展名）作为title
        file_type = os.path.splitext(filename)[1][1:] or "txt"
        title = os.path.splitext(filename)[0] if os.path.splitext(filename)[0] else filename

        document = Document(
            user_id=user_id,
            title=title,
            file_type=file_type,
            file_size=file_size,
            filename=filename,
            file_path=file_path,
            mime_type=mime_type,
            file_hash=file_hash,
            status="processing"
        )

        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)

        # 异步处理文档（解析+索引）
        # 这里简化为同步处理
        await self._process_document(document)

        return document

    async def _process_document(self, document: Document):
        """处理文档（解析+索引）"""
        try:
            # 1. 解析文档内容
            text = await self._parse_document(document.file_path, document.mime_type)

            if not text:
                document.status = "error"
                document.error_message = "文档内容为空"
                await self.db.commit()
                return

            # 2. 更新文档信息
            document.total_chars = len(text)
            document.processed_at = datetime.utcnow()

            # 3. 文本分块
            chunks = self._chunk_text(text)

            # 4. 索引到向量库
            chunk_count = await self.rag_service.index_document(
                document_id=document.id,
                user_id=document.user_id,
                file_name=document.filename,
                chunks=chunks
            )

            # 5. 更新状态
            document.status = "indexed"
            document.chunk_count = chunk_count

            # 保存摘要
            document.summary = self._generate_summary(text)

            await self.db.commit()

            print(f"✅ 文档处理成功: {document.filename} ({chunk_count} chunks)")

        except Exception as e:
            document.status = "error"
            document.error_message = str(e)
            await self.db.commit()
            print(f"❌ 文档处理失败: {document.filename} - {e}")

    async def _parse_document(self, file_path: str, mime_type: str) -> str:
        """解析文档内容"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return await self._parse_pdf(file_path)
        elif ext == ".docx":
            return await self._parse_docx(file_path)
        elif ext in [".txt", ".md"]:
            return await self._parse_text_file(file_path)
        elif ext in [".xlsx", ".xls"]:
            return await self._parse_excel(file_path)
        else:
            return ""

    async def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文档"""
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            print(f"PDF解析失败: {e}")

        return text

    async def _parse_docx(self, file_path: str) -> str:
        """解析Word文档"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"DOCX解析失败: {e}")

        return text

    async def _parse_text_file(self, file_path: str) -> str:
        """解析文本文件（支持中文）"""
        text = ""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                text = await f.read()

            # 如果编码错误，尝试gbk
            if not text:
                async with aiofiles.open(file_path, "r", encoding="gbk") as f:
                    text = await f.read()
        except Exception as e:
            print(f"文本文件解析失败: {e}")

        return text

    async def _parse_excel(self, file_path: str) -> str:
        """解析Excel文档"""
        text = ""
        try:
            wb = load_workbook(file_path, read_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"--- Sheet: {sheet_name} ---\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) if cell is not None else "" for cell in row])
                    text += row_text + "\n"
        except Exception as e:
            print(f"Excel解析失败: {e}")

        return text

    def _chunk_text(self, text: str) -> List[str]:
        """智能文本分块"""
        chunks = []

        # 按段落分割
        paragraphs = text.split("\n\n")

        current_chunk = ""
        current_size = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_size = len(para)

            # 如果单独一个段落就超过chunk大小，按句子分割
            if para_size > self.chunk_size:
                # 保存当前chunk
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                    current_size = 0

                # 按句子分割
                sentences = para.split("。")
                for sent in sentences:
                    sent = sent.strip()
                    if not sent:
                        continue

                    if current_size + len(sent) > self.chunk_size:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = sent + "。"
                        current_size = len(current_chunk)
                    else:
                        current_chunk += sent + "。"
                        current_size += len(sent)
            else:
                # 段落可以加入当前chunk
                if current_size + para_size > self.chunk_size:
                    # 保存当前chunk
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = para
                    current_size = para_size
                else:
                    current_chunk += "\n" + para + "\n"
                    current_size += para_size

        # 保存最后一个chunk
        if current_chunk:
            chunks.append(current_chunk.strip())

        # 确保每个chunk不过小（除了最后一个）
        min_chunk_size = 50
        final_chunks = []
        for i, chunk in enumerate(chunks):
            if len(chunk) < min_chunk_size and i > 0:
                # 合并到前一个chunk
                final_chunks[-1] += "\n" + chunk
            else:
                final_chunks.append(chunk)

        return final_chunks

    def _generate_summary(self, text: str, max_length: int = 200) -> str:
        """生成文档摘要"""
        if len(text) <= max_length:
            return text

        return text[:max_length] + "..."

    async def _calculate_file_hash(self, file_path: str) -> str:
        """计算文件SHA256哈希"""
        sha256_hash = hashlib.sha256()
        async with aiofiles.open(file_path, "rb") as f:
            while True:
                byte_block = await f.read(8192)
                if not byte_block:
                    break
                sha256_hash.update(byte_block)

        return sha256_hash.hexdigest()

    async def get_user_documents(
        self,
        user_id: int,
        status: str = None,
        limit: int = 100
    ) -> List[Document]:
        """获取用户文档列表"""
        query = select(Document).where(Document.user_id == user_id)

        if status:
            query = query.where(Document.status == status)

        query = query.order_by(Document.created_at.desc()).limit(limit)

        result = await self.db.execute(query)
        documents = result.scalars().all()

        return documents

    async def delete_document(self, document_id: int, user_id: int) -> bool:
        """删除文档"""
        query = select(Document).where(
            Document.id == document_id,
            Document.user_id == user_id
        )
        result = await self.db.execute(query)
        document = result.scalar_one_or_none()

        if not document:
            return False

        # 删除向量库中的数据
        await self.rag_service.milvus.delete_by_document(document_id)

        # 删除文件
        if os.path.exists(document.file_path):
            os.remove(document.file_path)

        # 删除记录
        await self.db.delete(document)
        await self.db.commit()

        return True

    async def get_document_stats(self, user_id: int) -> Dict:
        """获取用户文档统计"""
        from sqlalchemy import func

        query = select(
            func.count(Document.id).label("total"),
            func.sum(Document.chunk_count).label("total_chunks"),
            func.sum(Document.file_size).label("total_size")
        ).where(Document.user_id == user_id)

        result = await self.db.execute(query)
        stats = result.first()

        return {
            "total_documents": stats.total or 0,
            "total_chunks": int(stats.total_chunks or 0),
            "total_size": int(stats.total_size or 0)
        }


# 需要在文件开头导入settings
from core.config import settings
